import io

from pypdf import PdfReader
from docx import Document


def clean_pdf_text(text):
    """
    Clean text extracted from PDFs.

    Handles PDFs where characters are extracted with spaces:
        R e a c t  -> React
        P y t h o n -> Python
        N o d e . j s -> Node.js
    """

    cleaned_lines = []

    for line in text.splitlines():
        line = line.strip()

        if not line:
            continue

        words = line.split()

        # Detect text where characters are separated by spaces.
        if len(words) >= 2:
            single_chars = sum(
                1 for word in words if len(word) == 1
            )

            if single_chars >= len(words) * 0.6:
                line = "".join(words)

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def extract_text_from_pdf(file_stream):
    """
    Extract text from a PDF resume.
    """

    reader = PdfReader(file_stream)

    text = ""

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return clean_pdf_text(text).strip()


def extract_text_from_docx(file_stream):
    """
    Extract text from a DOCX resume.

    Extracts both:
    - Normal paragraphs
    - Tables
    """

    document = Document(file_stream)

    text_parts = []

    # Extract normal paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                text_parts.append(" ".join(row_text))

    return "\n".join(text_parts).strip()


def extract_resume_text(file):
    """
    Detect uploaded resume format and extract text.

    Supported:
        PDF
        DOCX
    """

    filename = file.filename.lower()

    # Read uploaded file into memory
    file_bytes = file.read()

    if not file_bytes:
        raise ValueError("Uploaded resume is empty.")

    file_stream = io.BytesIO(file_bytes)

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_stream)

    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_stream)

    else:
        raise ValueError(
            "Unsupported file type. Only PDF and DOCX are allowed."
        )